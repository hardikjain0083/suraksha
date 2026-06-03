import re
import asyncio
import io
import time
import hashlib
import logging
from dataclasses import dataclass
from typing import List, Optional

from PyPDF2 import PdfReader
import fitz  # PyMuPDF
import pdfplumber
from docx import Document

from models.circular import Clause
from services.notification import send_notification
from config import settings
from services.ocr_extractor import OCRExtractor, OCRException, LowConfidenceException

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────
#  Embedder (lazy-loaded)
# ─────────────────────────────────────────────────────────────

_embedder = None
_use_mock = False


def get_embedder():
    global _embedder, _use_mock
    if _embedder is None and not _use_mock:
        try:
            from sentence_transformers import SentenceTransformer
            _embedder = SentenceTransformer("all-MiniLM-L6-v2", local_files_only=True)
            logger.info("SentenceTransformer loaded successfully.")
        except Exception as e:
            logger.warning(f"Failed to load local sentence_transformers model: {e}. Using deterministic fallback embeddings.")
            _use_mock = True
    return _embedder


async def generate_embeddings(texts: List[str]) -> List[List[float]]:
    if not texts:
        return []
    embedder = get_embedder()
    if embedder is None:
        embeddings = []
        for text in texts:
            digest = hashlib.sha256(text.encode("utf-8")).digest()
            embeddings.append([((digest[i % len(digest)] / 255.0) * 2) - 1 for i in range(384)])
        return embeddings
    loop = asyncio.get_event_loop()
    embeddings = await loop.run_in_executor(None, embedder.encode, texts)
    return embeddings.tolist()


# ─────────────────────────────────────────────────────────────
#  PDF Result dataclass
# ─────────────────────────────────────────────────────────────

@dataclass
class PDFExtractionResult:
    text: str
    tables: List[Any]  # Keep generic to support both nested lists and structured Textract tables
    is_index_page: bool
    confidence: float
    engine_used: str = "unknown"
    ocr_metadata: Optional[dict] = None
    flagged_for_manual_review: bool = False


# ─────────────────────────────────────────────────────────────
#  Helper functions — MUST be defined BEFORE they are called
# ─────────────────────────────────────────────────────────────

def _detect_index_page(text: str) -> bool:
    """Detect if this is an index/listing page rather than a regulatory circular."""
    indicators = [
        r"INDEX TO RBI CIRCULARS",
        r"INDEX TO",
        r"Circular Number.*Date Of Issue.*Department.*Subject",
        r"Circular\s+Number\s+Date\s+Of\s+Issue",
        r"Master Directions",
        r"Master Circulars",
        r"Notifications\s+.*\d{4}",
        r"Back to previous page",
        r"FOLLOW RBI",
        r"MORE LINKS",
        r"Ref\.No\.",
        r"Subject\s+Meant\s+For",
    ]
    score = sum(1 for pattern in indicators if re.search(pattern, text, re.IGNORECASE))
    return score >= 2


def _score_extraction(text: str) -> float:
    """Score extraction quality (0.0–1.0)."""
    if not text or len(text.strip()) < 100:
        return 0.0
    score = 0.5
    readable_ratio = len(re.findall(r"[a-zA-Z\s]{3,}", text)) / max(len(text), 1)
    score += readable_ratio * 0.3
    if re.search(r"\d+\.\d+", text):
        score += 0.1
    if re.search(r"(shall|must|should|may)", text, re.IGNORECASE):
        score += 0.1
    return min(score, 1.0)


def _is_noise_paragraph(text: str) -> bool:
    """Return True if the paragraph is website noise (nav, footer, etc.)."""
    noise_patterns = [
        r"^Home\s+About Us\s+Notifications",
        r"^Master Directions\s+Master Circulars",
        r"^FOLLOW RBI",
        r"^Bank Holidays\s+Contact Us",
        r"^\d+\s+of\s+\d+$",
        r"^Download Mobile App",
        r"^Play Store\s+App Store",
    ]
    return any(re.search(p, text.strip(), re.IGNORECASE) for p in noise_patterns)


def _clean_rbi_text(text: str) -> str:
    """Remove RBI website header/footer noise from extracted text."""
    # Remove bilingual RBI branding header
    text = re.sub(
        r"भारतीय रिज़र्व बैंक.*?India's Central Bank", "", text, flags=re.DOTALL
    )
    text = re.sub(
        r"Reserve Bank of India.*?India's Central Bank", "", text, flags=re.DOTALL
    )
    # Remove navigation
    text = re.sub(
        r"Home\s+About Us\s+Notifications.*?Regulatory Reporting", "", text, flags=re.DOTALL
    )
    # Remove footer
    text = re.sub(r"Back to previous page.*", "", text, flags=re.DOTALL)
    text = re.sub(r"MORE LINKS.*?App Store", "", text, flags=re.DOTALL)
    text = re.sub(r"RBI's Vision and Values.*?Tenders", "", text, flags=re.DOTALL)
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ─────────────────────────────────────────────────────────────
#  Multi-engine PDF extraction — helpers are now defined above
# ─────────────────────────────────────────────────────────────

async def extract_pdf_robust(file_bytes: bytes) -> PDFExtractionResult:
    """Try PyMuPDF → pdfplumber → PyPDF2; pick best quality result."""
    results = []

    # Strategy 1: PyMuPDF (fitz) — best for text-heavy PDFs
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        raw = ""
        for page in doc:
            raw += page.get_text()
        doc.close()
        text = _clean_rbi_text(raw)
        is_index = _detect_index_page(text)
        quality = _score_extraction(text)
        results.append(("pymupdf", text, [], is_index, quality))
    except Exception as e:
        logger.error(f"PyMuPDF failed: {e}")

    # Strategy 2: pdfplumber — best when tables are present
    try:
        tables: List[List[List[str]]] = []
        raw = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                raw += page_text + "\n"
                for table in page.extract_tables():
                    if table and len(table) > 1:
                        tables.append(table)
        text = _clean_rbi_text(raw)
        is_index = _detect_index_page(text) or len(tables) > 0
        quality = min(_score_extraction(text) * (1.2 if tables else 1.0), 1.0)
        results.append(("pdfplumber", text, tables, is_index, quality))
    except Exception as e:
        logger.error(f"pdfplumber failed: {e}")

    # Strategy 3: PyPDF2 — fallback
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        raw = ""
        for page in reader.pages:
            raw += page.extract_text() or ""
        text = _clean_rbi_text(raw)
        is_index = _detect_index_page(text)
        quality = _score_extraction(text) * 0.8
        results.append(("pypdf2", text, [], is_index, quality))
    except Exception as e:
        logger.error(f"PyPDF2 failed: {e}")

    if not results:
        logger.error("All PDF extraction strategies failed")
        return PDFExtractionResult("", [], False, 0.0, "none")

    best = max(results, key=lambda x: x[4])
    engine, text, tables, is_index, quality = best
    logger.info(f"PDF extraction: engine={engine} quality={quality:.2f} is_index={is_index}")

    # Determine if we need to fall back to OCR
    ocr_extractor = OCRExtractor()
    is_scanned = ocr_extractor.is_scanned_pdf(file_bytes)
    
    # If the text is empty/garbled or is heuristically detected as scanned, trigger OCR
    if is_scanned or quality < settings.ocr_tesseract_min_confidence or len(text.strip()) < 100:
        logger.info(f"Scanned PDF or low quality detected (is_scanned={is_scanned}, quality={quality:.2f}, text_len={len(text.strip())}). Triggering OCR Fallback...")
        try:
            ocr_result = ocr_extractor.extract_text(file_bytes)
            return PDFExtractionResult(
                text=ocr_result["text"],
                tables=ocr_result["tables"],
                is_index_page=is_index or _detect_index_page(ocr_result["text"]),
                confidence=ocr_result["confidence"],
                engine_used=ocr_result["engine_used"],
                ocr_metadata=ocr_result,
                flagged_for_manual_review=False
            )
        except LowConfidenceException as e:
            logger.warning(f"OCR succeeded but confidence was below threshold: {e}")
            ocr_res = e.ocr_result
            return PDFExtractionResult(
                text=ocr_res.get("text", ""),
                tables=ocr_res.get("tables", []),
                is_index_page=is_index or _detect_index_page(ocr_res.get("text", "")),
                confidence=ocr_res.get("confidence", 0.0),
                engine_used=ocr_res.get("engine_used", "none"),
                ocr_metadata=ocr_res,
                flagged_for_manual_review=True
            )
        except OCRException as e:
            logger.error(f"OCR Extraction Exception: {e}")
            return PDFExtractionResult(
                text=text,
                tables=tables,
                is_index_page=is_index,
                confidence=quality,
                engine_used="failed_ocr",
                ocr_metadata={
                    "status": "failed",
                    "error": str(e),
                    "engine_used": settings.ocr_engine
                },
                flagged_for_manual_review=True
            )
        except Exception as e:
            logger.error(f"Unexpected error during OCR: {e}")
            return PDFExtractionResult(
                text=text,
                tables=tables,
                is_index_page=is_index,
                confidence=quality,
                engine_used="failed_ocr",
                ocr_metadata={
                    "status": "failed",
                    "error": str(e),
                    "engine_used": settings.ocr_engine
                },
                flagged_for_manual_review=True
            )

    return PDFExtractionResult(text, tables, is_index, quality, engine)


# ─────────────────────────────────────────────────────────────
#  Clause parsers
# ─────────────────────────────────────────────────────────────

def parse_clauses_from_index(tables: List[List[List[str]]]) -> List[Clause]:
    """Extract index entries from table rows."""
    clauses: List[Clause] = []
    for table in tables:
        if not table or len(table) < 2:
            continue

        headers = [str(h).lower().strip() for h in (table[0] or []) if h]
        col_num = next((i for i, h in enumerate(headers) if "number" in h or "circular" in h), 0)
        col_subj = next(
            (i for i, h in enumerate(headers) if "subject" in h or "topic" in h),
            min(3, len(table[0]) - 1),
        )

        for row in table[1:]:
            if not row or len(row) <= max(col_num, col_subj):
                continue
            c_num = re.sub(r"\s+", " ", str(row[col_num] or "").strip())
            c_text = str(row[col_subj] or "").strip()
            if not c_num and not c_text:
                continue
            clauses.append(
                Clause(
                    clause_number=c_num or None,
                    text=c_text or c_num,
                    obligation_type="shall",
                    severity="medium",
                    gap_status="pending",
                )
            )
    return clauses


def parse_index_from_text(text: str) -> List[Clause]:
    """Regex fallback: extract RBI index entries from plain text."""
    clauses: List[Clause] = []
    pattern = re.compile(
        r"(RBI/\d{4}-\d{2,4}/\d+.*?)(?=(?:RBI/|\Z))", re.DOTALL
    )
    for match in pattern.finditer(text):
        raw = match.group(1).strip()
        lines = [l.strip() for l in raw.split("\n") if l.strip()]
        if not lines:
            continue
        c_num = lines[0]
        c_text = " ".join(lines[1:])[:500] if len(lines) > 1 else c_num
        clauses.append(
            Clause(
                clause_number=c_num,
                text=c_text,
                obligation_type="shall",
                severity="medium",
                gap_status="pending",
            )
        )
    return clauses


def parse_clauses(text: str) -> List[Clause]:
    """Parse regulatory obligations from circular text."""
    clauses: List[Clause] = []
    
    # Matches: 1., 1.1, 1.1.1, (a), (i)
    clause_pattern = re.compile(r"^((?:\d+\.)+(?:\d+)?|\d+\.|\([a-zA-Z0-9]{1,3}\))\s+")
    
    # First, split into lines and combine lines that don't look like new clauses
    # into a single paragraph to avoid skipping short leading lines.
    raw_lines = text.split("\n")
    paragraphs = []
    current_para = ""
    
    for line in raw_lines:
        line = line.strip()
        if not line or _is_noise_paragraph(line):
            if current_para:
                paragraphs.append(current_para)
                current_para = ""
            continue
            
        # If it looks like a new clause, or it's empty, start a new paragraph
        if clause_pattern.match(line):
            if current_para:
                paragraphs.append(current_para)
            current_para = line
        else:
            if current_para:
                current_para += " " + line
            else:
                current_para = line
                
    if current_para:
        paragraphs.append(current_para)

    for para in paragraphs:
        if len(para) < 15:
            continue

        match = clause_pattern.search(para)
        clause_num = match.group(1) if match else None

        lower = para.lower()
        obligation: Optional[str] = None
        severity: Optional[str] = None

        if "shall " in lower:
            obligation, severity = "shall", "critical"
        elif "must " in lower:
            obligation, severity = "must", "critical"
        elif "mandatory" in lower or "required" in lower:
            obligation, severity = "must", "critical"
        elif "should " in lower:
            obligation, severity = "should", "high"
        elif "may " in lower:
            obligation, severity = "may", "medium"
        elif "recommended" in lower:
            obligation, severity = "recommended", "low"

        penalty_ref = None
        if re.search(r"\b(penalty|section|fine|liable|contravention)\b", lower):
            penalty_ref = "Detected potential penalty reference"

        if clause_num or obligation:
            clauses.append(
                Clause(
                    clause_number=clause_num,
                    text=para,
                    obligation_type=obligation,
                    severity=severity,
                    penalty_reference=penalty_ref,
                    gap_status="pending",
                )
            )

    return clauses


# ─────────────────────────────────────────────────────────────
#  Main processing pipeline
# ─────────────────────────────────────────────────────────────

async def process_circular(
    file_bytes: bytes, filename: str
) -> tuple:
    """Extract, parse, embed. Returns (status, clauses, duration_ms, confidence, full_text, ocr_metadata, flagged_for_manual_review)."""
    start = time.time()
    ext = filename.rsplit(".", 1)[-1].lower()

    text = ""
    is_index = False
    tables: List[Any] = []
    confidence = 1.0
    ocr_metadata = None
    flagged_for_manual_review = False

    if ext == "pdf":
        res = await extract_pdf_robust(file_bytes)
        text = res.text
        is_index = res.is_index_page
        tables = res.tables
        confidence = res.confidence
        ocr_metadata = res.ocr_metadata
        flagged_for_manual_review = res.flagged_for_manual_review
    elif ext == "docx":
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    else:
        # Plain text
        text = file_bytes.decode("utf-8", errors="ignore")

    # ── Parse clauses ─────────────────────────────────────────
    if is_index:
        clauses = parse_clauses_from_index(tables) if tables else []
        if not clauses:
            clauses = parse_index_from_text(text)
        if not clauses:
            clauses = parse_clauses(text)
    else:
        clauses = parse_clauses(text)

    # ── Determine ingestion status ────────────────────────────
    if not clauses:
        status = "failed"
    elif is_index:
        status = "fully_parsed"
    else:
        missing_numbers = sum(1 for c in clauses if not c.clause_number)
        missing_obs = sum(1 for c in clauses if not c.obligation_type)
        if missing_numbers == 0 and missing_obs == 0:
            status = "fully_parsed"
        elif missing_numbers > len(clauses) * 0.5:
            status = "failed"
        else:
            status = "partially_parsed"

    # ── Generate embeddings ───────────────────────────────────
    if clauses:
        texts = [c.text for c in clauses]
        embs = await generate_embeddings(texts)
        for c, emb in zip(clauses, embs):
            c.embedding = emb

    if status == "failed":
        await send_notification(
            user_id="EMP-INFOSEC-001",
            subject="Parsing Failed Alert",
            message=f"Failed to parse circular: {filename}. Please check system logs."
        )

    duration_ms = int((time.time() - start) * 1000)
    return status, clauses, duration_ms, confidence, text, ocr_metadata, flagged_for_manual_review
