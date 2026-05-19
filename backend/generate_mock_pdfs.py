import os
from fpdf import FPDF
from docx import Document

MOCK_DIR = "mock-data/circulars"
os.makedirs(MOCK_DIR, exist_ok=True)

def create_rbi_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(200, 10, text="Reserve Bank of India - Master Direction", align='C')
    pdf.ln(10)
    pdf.cell(200, 10, text="Ref: RBI/2026-045", align='C')
    pdf.ln(10)
    
    clauses = [
        "1.1 This Master Direction establishes baseline cybersecurity controls.",
        "2.1 Applicability extends to all Commercial Banks and NBFCs.",
        "3.1 All regulated entities shall implement multi-factor authentication for all privileged access by June 30, 2026.",
        "3.2 Entities should establish a dedicated SOC operational 24x7.",
        "4.1 Endpoint detection and response (EDR) agents must be deployed on all endpoints.",
        "4.2 It is recommended to perform adversary simulation twice a year.",
        "5.1 Banks must conduct quarterly vulnerability assessments and maintain records for 7 years.",
        "6.1 Failure to comply may result in supervisory action under Section 35A."
    ]
    
    for clause in clauses:
        pdf.multi_cell(190, 10, text=clause.strip())
        pdf.ln(2)
        
    pdf.output(os.path.join(MOCK_DIR, "rbi-cybersecurity-2026-045.pdf"))
    print("Generated RBI PDF")

def create_sebi_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(200, 10, text="Securities and Exchange Board of India", align='C')
    pdf.ln(10)
    pdf.cell(200, 10, text="Circular: SEBI/2026-012", align='C')
    pdf.ln(10)
    
    clauses = [
        "1.1 Brokers shall secure their trading networks.",
        "2.1 All trading APIs must use TLS 1.3.",
        "Data at rest encryption is mandatory.",
        "3.1 You should rotate passwords every 90 days.",
        "4.1 Incident response playbooks shall be updated annually.",
        "Vendors must be audited.",
        "5.1 Brokers may deploy AI for fraud detection.",
        "6.1 Logs shall be retained for 5 years.",
        "7.1 Network segmentation must be strictly enforced.",
        "8.1 Penetration testing should be conducted bi-annually.",
        "9.1 Third-party risks shall be assessed.",
        "10.1 Access controls must be reviewed."
    ]
    
    for clause in clauses:
        pdf.multi_cell(190, 10, text=clause.strip())
        pdf.ln(5)
        
    pdf.output(os.path.join(MOCK_DIR, "sebi-broker-cyber-2026-012.pdf"))
    print("Generated SEBI PDF")

def create_cert_in_pdf():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=12)
    pdf.cell(200, 10, text="CERT-In Advisory 2026-003", align='C')
    pdf.ln(10)
    
    pdf.multi_cell(190, 10, text="PNG IHDRcIDATxk0;")
    pdf.multi_cell(190, 10, text="1.1 Organizations shall")
    pdf.multi_cell(190, 10, text="2.1   must")
    
    pdf.output(os.path.join(MOCK_DIR, "cert-in-failed-2026-003.pdf"))
    print("Generated CERT-In PDF")

def create_internal_docx():
    doc = Document()
    doc.add_heading('Internal Memo: Zero Trust Architecture', 0)
    doc.add_paragraph('1.1 All internal microservices shall verify JWT tokens on every request.')
    doc.add_paragraph('2.1 Developers must not commit secrets to Git repositories.')
    doc.add_paragraph('3.1 It is recommended to use HashiCorp Vault.')
    doc.save(os.path.join(MOCK_DIR, "internal-memo.docx"))
    print("Generated Internal DOCX")

if __name__ == "__main__":
    create_rbi_pdf()
    create_sebi_pdf()
    create_cert_in_pdf()
    create_internal_docx()
